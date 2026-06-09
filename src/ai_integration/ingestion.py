import os
import sys
import pdfplumber
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Import modular components
import config
from em_model import embeddings_model
import vector_store_manager as vs_manager

# Resolve Knowledge Base Dir
KNOWLEDGE_BASE_DIR = os.path.abspath(os.path.join(config.AI_DIR, "..", "..", "knowledge_base"))

def load_documents(directory_path: str):
    """
    Loads text, md, pdf, and docx documents from a directory.
    """
    documents = []
    if not os.path.exists(directory_path):
        print(f"[Ingestion] Directory {directory_path} does not exist.")
        return documents

    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        if os.path.isdir(file_path):
            continue
        
        ext = os.path.splitext(filename)[1].lower()
        content = ""
        metadata = {"source": filename}

        try:
            if ext in [".txt", ".md"]:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            elif ext == ".pdf":
                with pdfplumber.open(file_path) as pdf:
                    pages_content = []
                    for page in pdf.pages:
                        page_text = page.extract_text(layout=True)
                        if page_text:
                            pages_content.append(page_text)
                    content = "\n".join(pages_content)
            elif ext == ".docx":
                doc = DocxDocument(file_path)
                paras = [para.text for para in doc.paragraphs]
                content = "\n".join(paras)
            else:
                continue

            if content.strip():
                documents.append(Document(page_content=content, metadata=metadata))
                print(f"[Ingestion] Successfully loaded: {filename}")
        except Exception as e:
            print(f"[Ingestion] Error loading {filename}: {e}")

    return documents

def run_ingestion():
    """
    Ingests files from the knowledge base directory and writes them to the FAISS database store.
    """
    print(f"[Ingestion] Starting document ingestion from: {KNOWLEDGE_BASE_DIR}")
    
    if not os.path.exists(KNOWLEDGE_BASE_DIR):
        os.makedirs(KNOWLEDGE_BASE_DIR, exist_ok=True)
        print(f"[Ingestion] Created empty knowledge base directory at: {KNOWLEDGE_BASE_DIR}")

    docs = load_documents(KNOWLEDGE_BASE_DIR)
    
    if not docs:
        print("[Ingestion] No documents found to ingest. Skipping indexing.")
        return False

    print(f"[Ingestion] Loaded {len(docs)} documents. Splitting into chunks...")
    
    # Split documents
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150,
        length_function=len
    )
    chunks = text_splitter.split_documents(docs)
    print(f"[Ingestion] Created {len(chunks)} text chunks.")

    # Initialize Embeddings model (SafeGoogleGenerativeAIEmbeddings with fallback)
    embeddings = embeddings_model()

    # Create and Save Vector Store
    try:
        vector_store = vs_manager.create_vector_store(chunks, embeddings)
        vs_manager.save_vector_store(vector_store, config.VECTOR_STORE_PATH)
        print("[Ingestion] Ingestion completed successfully!")
        return True
    except Exception as e:
        print(f"[Ingestion] Error creating FAISS index: {e}")
        return False

if __name__ == "__main__":
    success = run_ingestion()
    sys.exit(0 if success else 1)
